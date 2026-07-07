"""Word-Bericht (.docx) aus KI-Analyseabschnitten."""

from __future__ import annotations

from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches


def _add_paragraphs(doc: Document, text: str) -> None:
    for para in text.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())


def build_run_report_docx(
    path: Path,
    *,
    title: str,
    subtitle: str,
    summary: str,
    text_sections: list[tuple[str, str]],
    image_sections: list[tuple[str, str, Path]],
    model_label: str,
    analysis_sections: list[tuple[str, str]] | None = None,
    appendix_sections: list[tuple[str, str]] | None = None,
    layout: str = "standard",
    numbered_sections: list[tuple[str, str]] | None = None,
    section_images: dict[str, list[tuple[str, str, Path]]] | None = None,
) -> Path:
    """Erstellt einen KI-Bericht (.docx) im Zielordner."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    doc.add_heading(title, level=0)
    doc.add_paragraph(subtitle)
    meta = doc.add_paragraph()
    meta.add_run(
        f"Erstellt: {datetime.now().strftime('%Y-%m-%d %H:%M')} · KI-Modell: {model_label}"
    ).italic = True
    doc.add_paragraph("")

    if layout == "tsa_model" and numbered_sections:
        images_by_section: dict[str, list[tuple[str, str, Path]]] = defaultdict(list)
        if section_images:
            for key, imgs in section_images.items():
                images_by_section[key].extend(imgs)
        for heading, body in numbered_sections:
            doc.add_heading(heading, level=1)
            _add_paragraphs(doc, body)
            for img_heading, explanation, img_path in images_by_section.get(heading, []):
                doc.add_heading(img_heading, level=2)
                if img_path.is_file():
                    try:
                        doc.add_picture(str(img_path), width=Inches(5.8))
                    except Exception:
                        doc.add_paragraph(
                            f"(Grafik konnte nicht eingebettet werden: {img_path.name})"
                        )
                _add_paragraphs(doc, explanation)
            doc.add_paragraph("")
        if appendix_sections:
            doc.add_heading("Appendix — Source files", level=1)
            for heading, body in appendix_sections:
                doc.add_heading(heading, level=2)
                _add_paragraphs(doc, body)
                doc.add_paragraph("")
    else:
        doc.add_heading("Executive Summary", level=1)
        _add_paragraphs(doc, summary)
        doc.add_paragraph("")

        analysis = analysis_sections if analysis_sections is not None else text_sections
        appendix = appendix_sections if appendix_sections is not None else []

        if analysis:
            doc.add_heading("Fachliche Auswertung", level=1)
            for heading, body in analysis:
                doc.add_heading(heading, level=2)
                _add_paragraphs(doc, body)
                doc.add_paragraph("")

        if appendix:
            doc.add_heading("Anhang — Ausgabedateien", level=1)
            for heading, body in appendix:
                doc.add_heading(heading, level=2)
                _add_paragraphs(doc, body)
                doc.add_paragraph("")

        if image_sections:
            doc.add_heading("Grafiken und Visualisierungen", level=1)
            for heading, explanation, img_path in image_sections:
                doc.add_heading(heading, level=2)
                if img_path.is_file():
                    try:
                        doc.add_picture(str(img_path), width=Inches(5.8))
                    except Exception:
                        doc.add_paragraph(
                            f"(Grafik konnte nicht eingebettet werden: {img_path.name})"
                        )
                _add_paragraphs(doc, explanation)
                doc.add_paragraph("")

    doc.save(path)
    return path
