"""Tests fuer Word-Bericht und Report-Service."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from tslab.services.report_docx import build_run_report_docx
from tslab.services.report_service import generate_run_report, load_report_config


class ReportDocxTests(unittest.TestCase):
    def test_build_docx_contains_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "ai_bericht.docx"
            build_run_report_docx(
                out,
                title="Test TSLab",
                subtitle="Ordner: test_run",
                summary="Kurze Executive Summary.",
                text_sections=[("stats.txt", "r = 0.85")],
                image_sections=[],
                model_label="OpenAI GPT-4o mini",
            )
            self.assertTrue(out.is_file())
            doc = Document(out)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Test TSLab", text)
            self.assertIn("Executive Summary", text)
            self.assertIn("stats.txt", text)
            self.assertIn("r = 0.85", text)

    def test_tsa_model_layout_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "TSA_Modell_Bericht_ARMA11_test.docx"
            build_run_report_docx(
                out,
                title="TSA Modell arma11",
                subtitle="Ordner: arma11",
                summary="",
                text_sections=[],
                image_sections=[],
                model_label="GPT-5 nano",
                layout="tsa_model",
                numbered_sections=[
                    ("1. Management Summary", "Summary text."),
                    ("2. Introduction of the TSA", "Intro text."),
                    ("6. Conclusion", "Done."),
                ],
            )
            doc = Document(out)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Management Summary", text)
            self.assertIn("Introduction of the TSA", text)
            self.assertIn("Conclusion", text)
            self.assertNotIn("Executive Summary", text)


class GenerateRunReportTests(unittest.TestCase):
    def test_disabled_returns_status(self) -> None:
        cfg = load_report_config()
        if cfg.enabled:
            self.skipTest("ai_reports enabled in environment")
        result = generate_run_report("output/nonexistent")
        self.assertFalse(result["ok"])
        self.assertEqual(result["status"], "disabled")

    def test_empty_folder_error(self) -> None:
        from tests.helpers_output import temp_output_run

        with temp_output_run("empty_report_run") as run:
            with patch("tslab.services.report_service.load_report_config") as mock_cfg:
                mock_cfg.return_value = MagicMock(
                    enabled=True,
                    max_tokens=100,
                    default_model="openai:gpt-4o-mini",
                    output_basename="ai_bericht.docx",
                    openai_api_key="test",
                    gemini_api_key=None,
                    langfuse_public_key=None,
                    langfuse_secret_key=None,
                    langfuse_host=None,
                    models=(),
                )
                result = generate_run_report(
                    run, model_id="openai:gpt-4o-mini"
                )
        self.assertFalse(result["ok"])
        self.assertIn("Keine Berichtsziele", result["message"])


if __name__ == "__main__":
    unittest.main()
